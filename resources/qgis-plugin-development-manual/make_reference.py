from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
TEAL = RGBColor(11, 79, 92)
ORANGE = RGBColor(231, 122, 34)
DARK = RGBColor(32, 46, 49)
MUTED = RGBColor(92, 112, 115)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.0)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, colour, before, after in [
    ("Title", 28, TEAL, 0, 16),
    ("Subtitle", 15, MUTED, 0, 10),
    ("Heading 1", 20, TEAL, 18, 8),
    ("Heading 2", 15, TEAL, 14, 6),
    ("Heading 3", 12, ORANGE, 10, 4),
    ("Heading 4", 11, TEAL, 8, 3),
]:
    style = doc.styles[name]
    style.font.name = "Aptos Display" if name in {"Title", "Heading 1", "Heading 2"} else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = name != "Subtitle"
    style.font.color.rgb = colour
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("Source Code", "Verbatim Char"):
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.font.name = "DejaVu Sans Mono"
        style.font.size = Pt(8.3)
        style.font.color.rgb = RGBColor(30, 45, 48)

header = section.header.paragraphs[0]
header.text = "DESENVOLVIMENTO DE PLUGINS QGIS COM PYTHON"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Jubílio Filiano Maússe  |  ")
run.font.size = Pt(8)
run.font.color.rgb = MUTED
begin = OxmlElement("w:fldChar")
begin.set(qn("w:fldCharType"), "begin")
instruction = OxmlElement("w:instrText")
instruction.set(qn("xml:space"), "preserve")
instruction.text = " PAGE "
end = OxmlElement("w:fldChar")
end.set(qn("w:fldCharType"), "end")
run._r.extend([begin, instruction, end])

doc.save(ROOT / "manual" / "reference.docx")
